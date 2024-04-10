import os
import docx
import enchant
from language_tool_python import LanguageTool

def spell_check(text):
    d = enchant.Dict("en_US")
    corrected_text = []
    for word in text.split():
        if not d.check(word):
            suggestions = d.suggest(word)
            if suggestions:
                corrected_text.append(suggestions[0])
            else:
                corrected_text.append(word)
        else:
            corrected_text.append(word)
    return " ".join(corrected_text)

def grammar_check(text, tool):
    return tool.correct(text)

def process_text(input_file, output_file):
    file_ext = os.path.splitext(input_file)[1].lower()

    # Read input text from the file
    if file_ext == ".txt":
        with open(input_file, "r", encoding="utf-8") as f:
            input_text = f.read()
    elif file_ext == ".docx":
        doc = docx.Document(input_file)
        input_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        print("Unsupported file format. Please provide a .txt or .docx file.")
        return

    # Spell check
    corrected_text = spell_check(input_text)

    # Grammar check
    tool = LanguageTool('en-US')
    corrected_text = grammar_check(corrected_text, tool)

    # Preserve ending punctuation
    if input_text[-1] in ['.', '!', '?']:
        corrected_text += input_text[-1]

    # Write the edited text to the output file with the same format as the input file
    if file_ext == ".txt":
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(corrected_text)
    elif file_ext == ".docx":
        doc = docx.Document()
        for paragraph in corrected_text.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(output_file)

# Example usage:
input_file = "input.txt"  # Input file path (can be .txt or .docx)
output_file = "output.txt"  # Output file path (will have the same format as input)
process_text(input_file, output_file)
print("Processing done.")
